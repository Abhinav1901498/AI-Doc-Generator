import cv2
import pytesseract
import torch
from docx import Document
from datetime import datetime
import google.generativeai as genai
from dotenv import load_dotenv
import sqlite3
import os
import pypandoc
from pymongo import MongoClient
from bson import ObjectId
from fer import FER
import pypandoc
pypandoc.download_pandoc()
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
print("🔍 Tesseract Version:", pytesseract.get_tesseract_version())

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise RuntimeError("❌ GEMINI_API_KEY not found in .env file")

genai.configure(api_key=api_key)
model = genai.GenerativeModel("")

client = MongoClient("mongodb+srv://abhinav3258:abhi3258@cluster0.h9z4jp3.mongodb.net/?appName=Cluster00")
db = client["ai_doc_generator"]
collection = db["documents"]
def convert_docx_to_format(input_file, output_format):
    if not os.path.exists(input_file):
        print("❌ File not found.")
        return None
    output_file = os.path.splitext(input_file)[0] + f".{output_format}"
    try:
        pypandoc.convert_file(
            input_file,
            to=output_format,
            outputfile=output_file,
            extra_args=['--standalone']
        )
        print(f"✅ File converted successfully: {output_file}")
        return output_file
    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        return None
def capture_image_from_webcam(save_path="captured_image.jpg"):
    print("📷 Opening webcam... Press SPACE to capture, ESC to exit.")
    cam = cv2.VideoCapture(0)
    if not cam.isOpened():
        print("❌ Cannot access webcam.")
        return None
    while True:
        ret, frame = cam.read()
        if not ret:
            break
        cv2.putText(frame, "Press SPACE to Capture | ESC to Exit", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)
        cv2.imshow("AI Document Generator - Capture Mode", frame)
        key = cv2.waitKey(1)
        if key % 256 == 27:
            break
        elif key % 256 == 32:
            cv2.imwrite(save_path, frame)
            print(f"✅ Image captured and saved as {save_path}")
            break
    cam.release()
    cv2.destroyAllWindows()
    return save_path if os.path.exists(save_path) else None

def extract_text_from_image(image_path):
    img = cv2.imread(image_path)
    if img is None:
        return ""
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    text = pytesseract.image_to_string(gray)
    return text.strip()

def generate_text(topic, style):
    style_prompts = {
        "Essay": f"Write a detailed essay on {topic} including introduction, body, and conclusion.",
        "Report": f"Write a formal report about {topic}, including objectives, data, and findings.",
        "Blog": f"Write an engaging and friendly blog post about {topic}. Use emojis and a casual tone.",
        "Story": f"Write a short creative story about {topic}. Make it imaginative.",
        "Summary": f"Summarize {topic} clearly and concisely.",
        "Poem": f"Write a short and beautiful poem about {topic}."
    }
    prompt = style_prompts.get(style, style_prompts["Essay"])
    response = model.generate_content([prompt])
    return response.text

def detect_objects(image_path):
    model_yolo = torch.hub.load('ultralytics/yolov5', 'yolov5s', verbose=False)
    results = model_yolo(image_path)
    detections = results.pandas().xyxy[0]['name'].tolist()
    return list(set(detections))

def generate_text_from_objects(objects):
    topic = ", ".join(objects)
    prompt = f"Write a descriptive report about an image containing {topic}. Include observations, context, and possible uses."
    response = model.generate_content([prompt])
    return topic, response.text

def detect_emotion_from_image(image_path):
    detector = FER(mtcnn=True)
    img = cv2.imread(image_path)
    if img is None:
        return None
    emotions = detector.detect_emotions(img)
    if not emotions:
        return None
    top_emotion = max(emotions[0]['emotions'], key=emotions[0]['emotions'].get)
    print(f"🧠 Detected Emotion: {top_emotion.capitalize()}")
    return top_emotion

def generate_text_from_mood(emotion):
    mood_prompts = {
        "happy": "Write a joyful and positive poem about a bright and cheerful day.",
        "sad": "Write a comforting paragraph about sadness and hope.",
        "angry": "Write a reflective note on dealing with anger and peace.",
        "surprise": "Write an exciting story full of unexpected twists.",
        "fear": "Write a short story showing how courage overcomes fear.",
        "disgust": "Write a thought-provoking article about finding beauty in imperfection.",
        "neutral": "Write a calm and balanced daily diary entry."
    }
    prompt = mood_prompts.get(emotion.lower(), "Write a creative piece about human emotions and feelings.")
    response = model.generate_content([prompt])
    return response.text

def create_docx(content, topic, style):
    doc = Document()
    doc.add_heading(f"{topic.title()} - {style}", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    doc.add_paragraph(content)
    filename = f"AI_{style}_{topic.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"📄 Document saved as {filename}")
    return filename

def convert_docx_to_format(input_file, output_format):
    output_file = os.path.splitext(input_file)[0] + f".{output_format}"
    pypandoc.convert_file(input_file, output_format, format='docx', outputfile=output_file, extra_args=['--standalone'])
    print(f"✅ File converted successfully: {output_file}")
    return output_file

def save_to_mongodb(topic, style, content, filename):
    doc_data = {"topic": topic, "style": style, "content": content, "filename": filename, "created_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    collection.insert_one(doc_data)
    print("✅ Document saved in MongoDB successfully.")

def save_mood_to_mongodb(emotion, content, filename):
    doc_data = {"type": "mood_document", "emotion": emotion, "content": content, "filename": filename, "created_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    collection.insert_one(doc_data)
    print("📦 Saved mood document to MongoDB successfully.")

def view_all_documents():
    docs = list(collection.find({}, {"topic": 1, "style": 1, "filename": 1, "created_at": 1, "emotion": 1}))
    if not docs:
        print("⚠️ No documents found in MongoDB.")
        return
    print("\n🗂️ Stored Documents:")
    print("-" * 80)
    for d in docs:
        print(f"🆔 ID: {d['_id']}")
        print(f"📘 Topic: {d.get('topic', 'N/A')}")
        print(f"📝 Style/Emotion: {d.get('style', d.get('emotion', 'N/A'))}")
        print(f"📄 File: {d['filename']}")
        print(f"🕒 Created At: {d['created_at']}")
        print("-" * 80)

def view_document_by_id():
    _id = input("Enter the MongoDB _id of the document: ").strip()
    try:
        doc = collection.find_one({"_id": ObjectId(_id)})
        if not doc:
            print("❌ Document not found.")
            return
        print("\n📄 Document Details:")
        print(f"📘 Topic/Emotion: {doc.get('topic', doc.get('emotion', 'N/A'))}")
        print(f"📝 Type: {doc.get('style', doc.get('type', 'N/A'))}")
        print(f"📄 File: {doc['filename']}")
        print(f"🕒 Created At: {doc['created_at']}")
        print("\n🧾 Content:\n" + "-"*60)
        print(doc["content"])
        print("-" * 60)
    except Exception as e:
        print(f"⚠️ Invalid ID or Error: {e}")

def delete_document_by_id():
    _id = input("Enter the MongoDB _id of the document to delete: ").strip()
    try:
        result = collection.delete_one({"_id": ObjectId(_id)})
        if result.deleted_count:
            print("✅ Document deleted successfully.")
        else:
            print("❌ Document not found.")
    except Exception as e:
        print(f"⚠️ Invalid ID or Error: {e}")

if __name__ == "__main__":
    print("🧠 === AI DOCUMENT GENERATOR + EMOTION DETECTION ===")
    while True:
        print("\nOptions:")
        print("1️⃣ Generate document from typed topic")
        print("2️⃣ Generate document from image text (OCR)")
        print("3️⃣ Generate AI Vision Report (YOLO + ML)")
        print("4️⃣ Capture from webcam and generate document")
        print("5️⃣ Detect emotion from image and create mood-based document")
        print("6️⃣ Convert existing DOCX to another format")
        print("7️⃣ View all stored documents")
        print("8️⃣ View document by ID")
        print("9️⃣ Delete document by ID")
        print("🔟 Exit")

        choice = input("\nEnter your choice: ").strip()

        if choice == "1":
            topic = input("Enter a topic: ").strip()
            style = input("Enter style (Essay, Report, Blog, Story, Summary, Poem): ").strip() or "Essay"
            text = generate_text(topic, style)
            filename = create_docx(text, topic, style)
            save_to_mongodb(topic, style, text, filename)

        elif choice == "2":
            path = input("Enter image file path: ").strip()
            if not os.path.exists(path):
                print("❌ File not found.")
                continue
            topic = extract_text_from_image(path)
            text = generate_text(topic, "Report")
            filename = create_docx(text, topic, "Report")
            save_to_mongodb(topic, "Report", text, filename)

        elif choice == "3":
            path = input("Enter image file path: ").strip()
            if not os.path.exists(path):
                print("❌ File not found.")
                continue
            objects = detect_objects(path)
            if not objects:
                print("⚠️ No objects detected.")
                continue
            topic, report = generate_text_from_objects(objects)
            filename = create_docx(report, topic, "AI Vision Report")
            save_to_mongodb(topic, "AI Vision Report", report, filename)

        elif choice == "4":
            captured_path = capture_image_from_webcam()
            if not captured_path:
                continue
            emotion_detector = FER(mtcnn=True)
            img = cv2.imread(captured_path)
            result = emotion_detector.detect_emotions(img)
            if result:
                emotions = result[0]["emotions"]
                dominant_emotion = max(emotions, key=emotions.get)
            else:
                dominant_emotion = "Unknown"
            print(f"\nDetected Emotion: {dominant_emotion}")
            print("\nPlease provide your details:")
            name = input("Enter your name: ")
            dob = input("Enter your Date of Birth (DD-MM-YYYY): ")
            intro = input("Write a short introduction about yourself: ")
            text = f"Name: {name}\nDOB: {dob}\nEmotion: {dominant_emotion}\nIntroduction: {intro}\n\n"
            mood_prompts = {
                "happy": f"Write a cheerful and inspiring message for {name}, who seems happy today.",
                "sad": f"Write a comforting note to {name}, who appears a bit sad, filled with hope and positivity.",
                "angry": f"Write a calm reflection for {name}, who feels angry, promoting peace and balance.",
                "surprise": f"Write an exciting short story about {name} experiencing a wonderful surprise.",
                "fear": f"Write a motivational piece for {name} about facing fear and finding courage.",
                "disgust": f"Write an insightful article for {name} on finding beauty in imperfection.",
                "neutral": f"Write a balanced daily reflection for {name}, who is calm and thoughtful."
            }
            prompt = mood_prompts.get(dominant_emotion.lower(), f"Write a creative note for {name} expressing their emotions and positivity.")
            response = model.generate_content([prompt])
            full_report = text + response.text
            topic = f"{name}_Emotion_Report"
            filename = create_docx(full_report, topic, "Emotion Report")
            save_to_mongodb(topic, "Emotion Report", full_report, filename)
            print("\n✅ Emotion-based report generated and saved successfully!")

        elif choice == "5":
            import tkinter as tk
            from tkinter import filedialog
            print("🖼️ Choose an image file for emotion detection...")
            root = tk.Tk()
            root.withdraw()
            root.attributes('-topmost', True)
            pictures_dir = os.path.join(os.path.expanduser("~"), "Pictures")
            if not os.path.exists(pictures_dir):
                os.makedirs(pictures_dir)
            path = filedialog.askopenfilename(parent=root, title="Select Image for Emotion Detection", initialdir=pictures_dir, filetypes=[("Image Files", "*.jpg *.jpeg *.png *.bmp *.tiff")])
            root.destroy()
            if not path:
                print("❌ No file selected.")
                continue
            image_name = os.path.splitext(os.path.basename(path))[0]
            parts = image_name.replace('-', '_').split('_')
            name = parts[0].capitalize() if len(parts) > 0 else "Unknown"
            dob = parts[1] if len(parts) > 1 and parts[1].isdigit() else "Not specified"
            if len(parts) > 2:
                intro = ' '.join(parts[2:]).replace('-', ' ')
            else:
                intro = input("🗒️ Enter a short introduction for this person: ").strip()
                if not intro:
                    intro = "No introduction provided."
            print(f"👤 Name: {name}")
            print(f"🎂 DOB: {dob}")
            print(f"🗒️ Intro: {intro}")
            emotion = detect_emotion_from_image(path)
            if not emotion:
                print("😕 No emotion detected.")
                continue
            text = generate_text_from_mood(emotion)
            final_text = f"👤 Name: {name}\n🎂 DOB: {dob}\n🗒️ Introduction: {intro}\n\n🧠 Detected Emotion: {emotion.capitalize()}\n\n{text}"
            filename = create_docx(final_text, f"{name}_{emotion}", "Mood Document")
            save_mood_to_mongodb(emotion, final_text, filename)

        elif choice == "6":
            input_file = input("Enter .docx file path to convert: ").strip()
            if not os.path.exists(input_file):
                print("❌ File not found.")
                continue
            print("Available formats: pdf, rtf, md, odt, html")
            fmt = input("Enter desired output format: ").strip().lower()
            valid_formats = ["pdf", "txt", "rtf", "md", "odt", "html"]
            if fmt not in valid_formats:
                print("❌ Invalid format selected.")
                continue
            convert_docx_to_format(input_file, fmt)

        elif choice == "7":
            view_all_documents()

        elif choice == "8":
            view_document_by_id()

        elif choice == "9":
            delete_document_by_id()

        elif choice == "10":
            print("👋 Exiting AI Document Generator.")
            break

        else:
            print("❌ Invalid choice. Try again.")
